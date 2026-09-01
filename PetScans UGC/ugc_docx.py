"""Shared .docx rendering helpers for PetScans creator script packs.

One creator per folder; each folder has a `build-<name>-pack.py` that imports
this module for layout and supplies its own copy. Keeps every pack looking the
same so a creator who works on two campaigns gets one consistent document.

Requires python-docx:  ../PetScans-Meta-Campaign/venv/bin/python
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BODY = "Calibri"
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
ACCENT = RGBColor(0xC1, 0x44, 0x2E)


# ------------------------------------------------------------------ atoms

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    tcPr.append(el)


def _runs(paragraph, text, size, color, bold_default=False, italic=False):
    """Split on ** and alternate bold."""
    bold = bold_default
    for part in text.split("**"):
        if part:
            run = paragraph.add_run(part)
            run.font.name = BODY
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.bold = bold
            run.italic = italic
        bold = not bold


def rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "D8D8D8")
    bdr.append(bottom)
    pPr.append(bdr)


def rich(doc, text, size=11, color=INK, before=0, after=6, indent=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    _runs(p, text, size, color, italic=italic)
    return p


def heading(doc, text, size=16, color=INK, before=16, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = BODY
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = True
    return p


def bullets(doc, items, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        _runs(p, it, size, INK)


# ------------------------------------------------------------- components

def shotlist(doc, rows):
    """rows: (time, visual, audio)"""
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, label in enumerate(("Time", "What we see", "What you say")):
        shade(hdr[i], "F2F2F2")
        run = hdr[i].paragraphs[0].add_run(label)
        run.font.name = BODY
        run.font.size = Pt(9.5)
        run.bold = True
        run.font.color.rgb = MUTED
    for time, visual, audio in rows:
        cells = t.add_row().cells
        for i, text in enumerate((time, visual, audio)):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            _runs(p, text, 10, INK)
    for row in t.rows:
        row.cells[0].width = Inches(0.75)
        row.cells[1].width = Inches(2.5)
        row.cells[2].width = Inches(3.35)
    return t


def vo(doc, beat, text):
    """A 60s timed beat: coloured cue line, then the words to say."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(beat)
    run.font.name = BODY
    run.font.size = Pt(9.5)
    run.bold = True
    run.font.color.rgb = ACCENT

    q = doc.add_paragraph()
    q.paragraph_format.space_after = Pt(4)
    q.paragraph_format.left_indent = Inches(0.15)
    _runs(q, text, 11.5, INK)


def callout(doc, text, fill="FAF6F0"):
    box = doc.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    cell = box.rows[0].cells[0]
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    _runs(p, text, 10, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return box


def script_header(doc, number, title, meta, onscreen):
    heading(doc, f"Script {number} — {title}", size=17, before=22, after=2)
    rich(doc, meta, size=10, color=MUTED, after=8)
    callout(doc, "What you'll see on screen: " + onscreen)


def cut_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = BODY
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = ACCENT


def table(doc, headers, rows, widths=None, bold_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, label in enumerate(headers):
        shade(hdr[i], "F2F2F2")
        run = hdr[i].paragraphs[0].add_run(label)
        run.font.name = BODY
        run.font.size = Pt(9.5)
        run.bold = True
        run.font.color.rgb = MUTED
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            _runs(p, text, 10, INK, bold_default=(i == bold_col))
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
    style = doc.styles["Normal"]
    style.font.name = BODY
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    return doc


def cover(doc, title, subtitle, blurb_lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.name = BODY
    run.font.size = Pt(30)
    run.bold = True
    run.font.color.rgb = INK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(subtitle)
    run.font.name = BODY
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT

    for i, line in enumerate(blurb_lines):
        rich(doc, line, after=14 if i == len(blurb_lines) - 1 else 4)
