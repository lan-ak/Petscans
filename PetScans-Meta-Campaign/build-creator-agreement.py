#!/usr/bin/env python3
"""Render the PetScans UGC Creator Agreement as .docx and .md.

Usage:  ./venv/bin/python build-creator-agreement.py

Contract text lives in contract_content.py. This file is layout only.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from contract_content import (COMPANY, CONTENT, RULE, SIGNATURE_BLOCKS,
                              USE_ANCHORS)

OUT_DIR = Path(__file__).parent
BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)

# --------------------------------------------------------------------------
# .docx rendering
# --------------------------------------------------------------------------

TOKEN = re.compile(r"(\*\*.+?\*\*|\[[^\[\]]+\])")


def add_runs(paragraph, text):
    """Render **bold** and [PLACEHOLDER] inline markup into a paragraph."""
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("[") and part.endswith("]"):
            run = paragraph.add_run(part)
            run.bold = True
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            run = paragraph.add_run(part)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE


def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def add_page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt in ("PAGE",):
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {txt} "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run = p.add_run()
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)


def build_docx(path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        add_page_number_footer(section)

    for item in CONTENT:
        kind = item[0]

        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(item[1])
            r.bold = True
            r.font.name = BODY_FONT
            r.font.size = Pt(18)

        elif kind == "subtitle":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(18)
            r = p.add_run(item[1])
            r.font.name = BODY_FONT
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        elif kind == "h1":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(item[1].upper())
            r.bold = True
            r.font.name = BODY_FONT
            r.font.size = Pt(12)

        elif kind == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(item[1])
            r.bold = True
            r.font.name = BODY_FONT
            r.font.size = Pt(11)

        elif kind in ("p", "c"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if kind == "c":
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.first_line_indent = Inches(-0.4)
            add_runs(p, item[1])

        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.75)
            p.paragraph_format.space_after = Pt(4)
            add_runs(p, item[1])

        elif kind == "note":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(item[1])
            r.italic = True
            r.bold = True
            r.font.name = BODY_FONT
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

        elif kind == "table":
            _, headers, rows = item
            titled = any(h for h in headers)
            t = doc.add_table(rows=1 if titled else 0, cols=len(headers))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(headers if titled else []):
                cell = t.rows[0].cells[i]
                cell.text = ""
                r = cell.paragraphs[0].add_run(h)
                r.bold = True
                r.font.name = BODY_FONT
                r.font.size = Pt(10)
                shade(cell, "EFEFEF")
            for row in rows:
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = ""
                    add_runs(cells[i].paragraphs[0], val)
                    for run in cells[i].paragraphs[0].runs:
                        run.font.size = Pt(10)
            doc.add_paragraph()

        elif kind == "pagebreak":
            doc.add_page_break()

        elif kind == "sig":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run("SIGNATURE PAGE")
            r.bold = True
            r.font.name = BODY_FONT
            r.font.size = Pt(12)

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(24)
            add_runs(p, f"This page forms part of the Creator Agreement between "
                        f"**{COMPANY}** and [CREATOR LEGAL NAME], dated [EFFECTIVE DATE]. "
                        f"By signing, both sides agree to it.")

            for heading, fields in SIGNATURE_BLOCKS:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(10)
                r = p.add_run(heading)
                r.bold = True
                r.font.name = BODY_FONT
                r.font.size = Pt(11)

                for label, anchor in fields:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.left_indent = Inches(0.2)
                    tab = " " * (13 - len(label))
                    r = p.add_run(f"{label}{tab}{RULE}")
                    r.font.name = BODY_FONT
                    r.font.size = Pt(11)
                    if USE_ANCHORS and anchor:
                        a = p.add_run(f"\\{anchor}\\")
                        a.font.name = BODY_FONT
                        a.font.size = Pt(6)
                        a.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                doc.add_paragraph()

    doc.save(path)


# --------------------------------------------------------------------------
# .md rendering (kept in sync with the .docx)
# --------------------------------------------------------------------------

def build_md(path):
    out = []
    for item in CONTENT:
        kind = item[0]
        if kind == "title":
            out.append(f"# {item[1]}\n")
        elif kind == "subtitle":
            out.append(f"**{item[1]}**\n\n---\n")
        elif kind == "h1":
            out.append(f"## {item[1]}\n")
        elif kind == "h2":
            out.append(f"### {item[1]}\n")
        elif kind in ("p", "c"):
            out.append(f"{item[1]}\n")
        elif kind == "bullet":
            out.append(f"- {item[1]}\n")
        elif kind == "note":
            out.append(f"> **{item[1]}**\n")
        elif kind == "table":
            _, headers, rows = item
            if any(h for h in headers):
                out.append("| " + " | ".join(headers) + " |")
            else:
                out.append("| " + " | ".join([" "] * len(headers)) + " |")
            out.append("|" + "|".join(["---"] * len(headers)) + "|")
            for row in rows:
                out.append("| " + " | ".join(c if c else " " for c in row) + " |")
            out.append("")
        elif kind == "pagebreak":
            out.append("\n---\n")
        elif kind == "sig":
            out.append(f"## SIGNATURE PAGE\n")
            out.append(f"This page forms part of the Creator Agreement between "
                       f"**{COMPANY}** and [CREATOR LEGAL NAME], dated [EFFECTIVE DATE]. "
                       f"By signing, both sides agree to it.\n")
            for heading, fields in SIGNATURE_BLOCKS:
                out.append(f"**{heading}**\n")
                for label, _ in fields:
                    out.append(f"{label}  {RULE}\n")
    path.write_text("\n".join(out) + "\n")


if __name__ == "__main__":
    docx_path = OUT_DIR / "PetScans-Creator-Agreement-US-FINAL.docx"
    md_path = OUT_DIR / "PetScans-Creator-Agreement-US-FINAL.md"
    build_docx(docx_path)
    build_md(md_path)
    print(f"wrote {docx_path}")
    print(f"wrote {md_path}")
